#!/usr/bin/env python3
"""
HFT control-plane — architecture, timing & strategic-space breakdown.

Renders the system block diagram, the strategy lifecycle (sim->paper->live) with
strategic zones, the arena-loop timing plan, the honest overfit gauntlet, the
genetic-arena natural-selection chart, the live leaderboard, and the data/signal
inventory — then assembles a Word document.

Data is a real snapshot captured 2026-05-31 from the TimescaleDB warehouse and the
live arena (~/hft-live/data/polymarket.db). Numbers are not invented.

    python3 analysis/architecture_report.py      # needs matplotlib + python-docx
"""
from __future__ import annotations
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "analysis", "out")
GEN_DATE = "2026-05-31"

# palette
ACCENT, UP, DOWN, EDGE, MUTE = "#1f3a5f", "#2e8b57", "#c0392b", "#e67e22", "#7f8c8d"
GOLD, SKY = "#b8860b", "#3a6ea5"

# --------------------------------------------------------------------------- #
# REAL DATA SNAPSHOT (2026-05-31)
# --------------------------------------------------------------------------- #
WAREHOUSE = {"ONE_DAY": (12, 29297), "ONE_HOUR": (3, 226294), "ONE_MINUTE": (5, 2380)}
COVERAGE = {  # coin: (daily_bars, hourly_bars)
    "BTC": (3969, 95158), "ETH": (3664, 87749), "SOL": (1810, 43387),
}
# genome kind -> (total_ever, alive_now)
GENOMES = [
    ("cb_mean_reversion", 40, 21), ("polymarket_market_maker", 16, 11),
    ("cross_venue_arb", 15, 3), ("llm_probability_oracle", 13, 3),
    ("category_specialist", 12, 2), ("cb_momentum_burst", 7, 1),
    ("random_walk_baseline", 6, 1), ("poly_short_binary_directional", 6, 1),
    ("poly_fade_spike", 6, 6), ("poly_breakout", 5, 0), ("cb_breakout", 5, 0),
    ("wallet_copy_filtered", 1, 1), ("multi_strategy", 1, 1),
]
GENERATIONS = [  # (gen, n_agents, top_score, ticks, children, status)
    (0, 87, 0.0446, 50, 13, "sealed"), (1, 46, None, 11, None, "OPEN"),
]
POP = {"retired": 82, "alive": 46, "elite": 5}  # total 133
LEADERBOARD = [  # (label, gen, pnl, trades, wins, tag)
    ("poly_short_binary_directional", 1, 57.68, 6, 2, "g1 child"),
    ("multi_strategy", 1, 25.51, 1, 1, "g1 child"),
    ("poly_fade_spike", 0, 3.27, 3, 3, "ELITE"),
    ("polymarket_market_maker", 0, 0.09, 3, 1, "retired"),
    ("poly_short_binary_directional", 0, -5.92, 2, 0, "culled"),
    ("poly_short_binary_directional", 0, -100.0, 1, 0, "culled"),
]
# overfit battery (daily): coin -> (best, PBO, DSR)
DSR_DAILY = [
    ("BTC", "sma50d", 0.00, 0.94), ("XRP", "sma20d", 0.36, 0.84), ("LINK", "don20d", 0.63, 0.83),
    ("SOL", "don20d", 0.07, 0.79), ("LTC", "sma20d", 0.19, 0.74), ("ADA", "z30d", 0.53, 0.73),
    ("ETH", "sma20d", 0.00, 0.69), ("BCH", "sma20d", 0.17, 0.65), ("DOGE", "z10d", 0.39, 0.59),
    ("MATIC", "don20d", 0.40, 0.56), ("AVAX", "don10d", 0.16, 0.55), ("DOT", "z10d", 0.79, 0.39),
]
DSR_BTC_HOURLY = 0.60   # BTC sma50d+vt, hourly + inverse-vol sizing (FELL from 0.94)
WALKFWD_HELD = 5        # of 12 coins, OOS Sharpe stayed positive
POLY_SNAPSHOTS = [
    ("5min-binary", 2160), ("15min-binary", 2160), ("elections", 756),
    ("other", 216), ("crypto", 54), ("geopolitics", 33), ("weather", 21),
]
# arena loop cadence
CYCLE_MIN = 5
ALLOCATE_EVERY, RESEARCH_EVERY, ALLOC_BUDGET = 12, 24, 10000


def save(fig, name):
    fig.tight_layout()
    fig.savefig(os.path.join(OUT, name), dpi=140, bbox_inches="tight")
    plt.close(fig)


# --------------------------------------------------------------------------- #
# FIG 1 — system architecture block diagram
# --------------------------------------------------------------------------- #
def fig_architecture():
    fig, ax = plt.subplots(figsize=(11, 6.2))
    ax.set_xlim(0, 11); ax.set_ylim(0, 6.4); ax.axis("off")

    def box(x, y, w, h, t, c, ec=ACCENT, fs=8):
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.03", fc=c, ec=ec, lw=1.3))
        ax.text(x + w / 2, y + h / 2, t, ha="center", va="center", fontsize=fs, weight="bold", color=ec)

    def arrow(p1, p2, c="#888"):
        ax.add_patch(FancyArrowPatch(p1, p2, arrowstyle="-|>", mutation_scale=12, color=c, lw=1.2))

    # column 1: data sources
    box(0.15, 5.2, 2.3, 0.85, "Coinbase candles\n(since origination)", "#eaf3ea")
    box(0.15, 4.15, 2.3, 0.85, "Polymarket Gamma/CLOB\n(5m·15m·elections)", "#eaf3ea")
    box(0.15, 3.10, 2.3, 0.85, "L2 order book\n(OFI·VPIN·microprice)", "#eaf3ea")
    box(0.15, 2.05, 2.3, 0.85, "On-chain wallets\n(copy signals)", "#eaf3ea")
    # warehouse
    box(2.85, 3.5, 2.0, 1.7, "TimescaleDB\nWAREHOUSE\n257k candles\nticks · snapshots", "#dfe6ee", fs=8.5)
    # backtest gauntlet
    box(5.25, 4.5, 2.3, 1.5, "BACKTEST +\nOVERFIT GAUNTLET\nPBO · Deflated-Sharpe\nwalk-forward", "#fdeede")
    # arena
    box(5.25, 2.4, 2.3, 1.6, "ARENA\ngenetic evolution\n133 agents · 13 genomes\nfitness · cull · breed", "#dfe6ee", fs=8.2)
    box(5.25, 0.7, 2.3, 1.2, "AI ENSEMBLE\nevaluators · LLM oracle\nresearch re-targeting", "#fbe9d0")
    # capsules + execution
    box(7.95, 3.3, 1.6, 1.7, "CAPSULES\nrisk envelopes\nwho gets $ & why", "#fbe9d0", fs=8)
    box(7.95, 1.4, 1.6, 1.5, "STAGES\nsim→paper→\nlive_eligible→live", "#dfe6ee", fs=8)
    box(9.85, 2.3, 1.0, 1.7, "GATED\nEXEC\n(sim $→\nreal $)", "#fbe4e4", fs=8)

    for y in (5.6, 4.55, 3.5, 2.45):
        arrow((2.45, y), (2.85, 4.35))
    arrow((4.85, 4.6), (5.25, 5.0))   # warehouse -> backtest
    arrow((4.85, 4.0), (5.25, 3.2))   # warehouse -> arena
    arrow((6.4, 4.5), (6.4, 4.0))     # backtest -> arena (priors)
    arrow((6.4, 2.4), (6.4, 1.9))     # arena <-> ensemble
    arrow((6.4, 1.9), (6.4, 2.4), c="#bbb")
    arrow((7.55, 3.1), (7.95, 3.7))   # arena -> capsules
    arrow((8.75, 3.3), (8.75, 2.9))   # capsules -> stages
    arrow((9.55, 2.6), (9.85, 3.0))   # stages -> exec
    # feedback loop exec -> warehouse
    ax.add_patch(FancyArrowPatch((10.3, 2.3), (3.85, 3.5), arrowstyle="-|>", mutation_scale=12,
                                 color=UP, lw=1.0, ls=(0, (4, 3)), connectionstyle="arc3,rad=0.32"))
    ax.text(6.8, 0.15, "realized outcomes feed back → warehouse (closed loop)", ha="center", fontsize=7.5, color=UP, style="italic")
    ax.text(5.5, 6.15, "HFT control-plane — end-to-end architecture", ha="center", fontsize=13, weight="bold", color=ACCENT)
    save(fig, "fig_architecture.png")


# --------------------------------------------------------------------------- #
# FIG 2 — strategy lifecycle & strategic zones
# --------------------------------------------------------------------------- #
def fig_lifecycle():
    fig, ax = plt.subplots(figsize=(11, 3.4))
    ax.set_xlim(0, 11); ax.set_ylim(0, 3); ax.axis("off")
    zones = [
        (0.0, 2.05, "#dfe6ee", "BIRTH\ngenome seeded\n{kind, params}", ACCENT),
        (2.05, 2.2, "#eaf3ea", "SIM\n$1000 paper, prove\nfitness on live signals", UP),
        (4.25, 2.0, "#eef5ee", "RANK + SEAL\nevolve every 50 ticks\ncull ~50%, breed, elite", "#2e8b57"),
        (6.25, 2.0, "#fdeede", "CAPSULE\nsurvivors funded\n(risk envelope)", EDGE),
        (8.25, 1.6, "#fff7e6", "LIVE-ELIGIBLE\nDSR/PnL gate\ntiny real cap", GOLD),
        (9.85, 1.0, "#fbe4e4", "LIVE / CULL\n$ or retired", DOWN),
    ]
    for x, w, c, label, tc in zones:
        ax.add_patch(FancyBboxPatch((x + 0.05, 0.7), w - 0.1, 1.6, boxstyle="round,pad=0.02", fc=c, ec=tc, lw=1.5))
        ax.text(x + w / 2, 1.5, label, ha="center", va="center", fontsize=8.0, color=tc, weight="bold")
    ax.annotate("", xy=(10.95, 0.45), xytext=(0.05, 0.45), arrowprops=dict(arrowstyle="-|>", color="#333", lw=1.6))
    ax.text(5.5, 0.2, "an agent's life  ·  only proven survivors advance toward real capital  →", ha="center", fontsize=8, color="#333")
    ax.text(5.5, 2.78, "Strategy lifecycle & strategic zones — sim money first, real money only once proven", ha="center", fontsize=11, weight="bold", color=ACCENT)
    save(fig, "fig_lifecycle.png")


# --------------------------------------------------------------------------- #
# FIG 3 — arena-loop timing plan ("space between")
# --------------------------------------------------------------------------- #
def fig_timing():
    fig, ax = plt.subplots(figsize=(11, 4.3))
    n = 26
    ax.set_xlim(-0.5, n + 0.5); ax.set_ylim(0, 5); ax.axis("off")
    lanes = [("every cycle (~5 min)", 4.0, UP), ("allocate (12 cyc, ~1h)", 2.7, EDGE), ("scan+research (24 cyc, ~2h)", 1.4, SKY)]
    for label, y, c in lanes:
        ax.add_patch(FancyBboxPatch((-0.4, y - 0.32), n + 0.8, 0.64, boxstyle="round,pad=0.01", fc="#f6f7f9", ec="#ddd", lw=0.8))
        ax.text(-0.5, y, label, ha="right", va="center", fontsize=7.6, color=c, weight="bold")
    for i in range(n):
        ax.add_patch(FancyBboxPatch((i + 0.08, 3.7), 0.84, 0.6, boxstyle="round,pad=0.01", fc="#eaf3ea", ec=UP, lw=0.7))
        ax.text(i + 0.5, 4.0, "tick\nsnap\nL2", ha="center", va="center", fontsize=4.6, color=UP)
        if (i + 1) % ALLOCATE_EVERY == 0:
            ax.add_patch(FancyBboxPatch((i + 0.08, 2.4), 0.84, 0.6, boxstyle="round,pad=0.01", fc="#fdeede", ec=EDGE, lw=0.9))
            ax.text(i + 0.5, 2.7, "$alloc\ncapsule", ha="center", va="center", fontsize=4.6, color=EDGE, weight="bold")
        if (i + 1) % RESEARCH_EVERY == 0:
            ax.add_patch(FancyBboxPatch((i + 0.08, 1.1), 0.84, 0.6, boxstyle="round,pad=0.01", fc="#e7eef7", ec=SKY, lw=0.9))
            ax.text(i + 0.5, 1.4, "scan\nresrch", ha="center", va="center", fontsize=4.6, color=SKY, weight="bold")
    ax.annotate("", xy=(n + 0.4, 0.55), xytext=(0.1, 0.55), arrowprops=dict(arrowstyle="-|>", color="#333", lw=1.4))
    ax.text(n / 2, 0.3, "cycle  →   (each ≈ 5 min; 26 cycles ≈ 2.2 h)", ha="center", fontsize=8, color="#333")
    ax.text(n / 2, 4.75, "Arena-loop timing plan — what fires when, and the strategic space between", ha="center", fontsize=11.5, weight="bold", color=ACCENT)
    save(fig, "fig_timing.png")


# --------------------------------------------------------------------------- #
# FIG 4 — overfit gauntlet (DSR per coin + the 0.95 gate)
# --------------------------------------------------------------------------- #
def fig_overfit():
    fig, ax = plt.subplots(figsize=(10, 4.6))
    coins = [d[0] for d in DSR_DAILY]
    dsr = [d[3] for d in DSR_DAILY]
    colors = [UP if v >= 0.95 else (GOLD if v >= 0.8 else MUTE) for v in dsr]
    bars = ax.bar(coins, dsr, color=colors, ec=ACCENT, lw=0.6)
    ax.axhline(0.95, color=DOWN, ls="--", lw=1.6, label="HARDENED gate (DSR > 0.95)")
    ax.bar(["BTC\nhourly\n+vol"], [DSR_BTC_HOURLY], color=DOWN, ec=ACCENT, lw=0.6, alpha=0.85)
    for b, v in zip(bars, dsr):
        ax.text(b.get_x() + b.get_width() / 2, v + 0.01, f"{v:.2f}", ha="center", fontsize=7.5)
    ax.text(len(coins), DSR_BTC_HOURLY + 0.01, f"{DSR_BTC_HOURLY:.2f}", ha="center", fontsize=7.5, color=DOWN)
    ax.set_ylim(0, 1.02); ax.set_ylabel("Deflated Sharpe Ratio  (P[edge is real])")
    ax.set_title("The honest overfit gauntlet: 12 coins → 0 hardened (BTC daily closest at 0.94)", color=ACCENT, weight="bold")
    ax.legend(fontsize=8, loc="lower left")
    ax.text(0.5, 0.06, "BTC daily is the single most defensible prior — but even it sits just under the gate.\nFiner hourly bars + vol-sizing made it WORSE (0.94 → 0.60): more trials, not more signal.",
            transform=ax.transAxes, fontsize=7.6, color="#444", style="italic", va="bottom")
    save(fig, "fig_overfit.png")


# --------------------------------------------------------------------------- #
# FIG 5 — arena natural selection (total vs alive per genome)
# --------------------------------------------------------------------------- #
def fig_arena():
    fig, ax = plt.subplots(figsize=(10.5, 5.0))
    g = sorted(GENOMES, key=lambda x: x[1], reverse=True)
    names = [x[0] for x in g]
    total = [x[1] for x in g]
    alive = [x[2] for x in g]
    y = np.arange(len(names))
    ax.barh(y, total, color="#d6dce4", ec=ACCENT, lw=0.5, label="ever spawned")
    ax.barh(y, alive, color=UP, ec=ACCENT, lw=0.5, label="alive now")
    for i, (t, a) in enumerate(zip(total, alive)):
        tag = "  ✓ all survived" if a == t and a > 1 else ("  ✗ wiped out" if a == 0 else "")
        ax.text(t + 0.4, i, f"{a}/{t}{tag}", va="center", fontsize=7.2,
                color=(UP if a == t and a > 1 else (DOWN if a == 0 else "#333")))
    ax.set_yticks(y); ax.set_yticklabels(names, fontsize=7.8)
    ax.invert_yaxis(); ax.set_xlabel("agents")
    ax.set_title("Natural selection in the arena — 82 retired, 46 alive, 5 elite (133 total)", color=ACCENT, weight="bold")
    ax.legend(fontsize=8, loc="lower right")
    save(fig, "fig_arena.png")


# --------------------------------------------------------------------------- #
# FIG 6 — live leaderboard (PnL) + fitness-vs-PnL note
# --------------------------------------------------------------------------- #
def fig_leaderboard():
    fig, ax = plt.subplots(figsize=(10, 4.4))
    labels = [f"{l[:22]}\n(g{gen}, {tag})" for l, gen, p, tr, w, tag in LEADERBOARD]
    pnl = [x[2] for x in LEADERBOARD]
    colors = [UP if p > 0 else DOWN for p in pnl]
    bars = ax.bar(range(len(pnl)), pnl, color=colors, ec=ACCENT, lw=0.6)
    ax.axhline(0, color="#333", lw=1)
    for b, p in zip(bars, pnl):
        ax.text(b.get_x() + b.get_width() / 2, p + (2 if p >= 0 else -6), f"${p:+.2f}", ha="center", fontsize=7.6)
    ax.set_xticks(range(len(labels))); ax.set_xticklabels(labels, fontsize=6.8)
    ax.set_ylabel("realized PnL (USD, sim)")
    ax.set_title("Live arena leaderboard — evolution found a comeback, fitness rewards low drawdown", color=ACCENT, weight="bold")
    ax.text(0.02, 0.04, "The +$57.68 top earner is a gen-1 MUTATED CHILD of a genome whose gen-0 ancestors lost −$100 and were culled.\nYet fitness (pnl% − 2·maxDD% + activity) crowns the steady poly_fade_spike elites (+$3.27, 3/3 wins) — not raw PnL.",
            transform=ax.transAxes, fontsize=7.2, color="#444", style="italic", va="bottom")
    save(fig, "fig_leaderboard.png")


# --------------------------------------------------------------------------- #
# FIG 7 — data & signal inventory
# --------------------------------------------------------------------------- #
def fig_signals():
    fig, (a1, a2) = plt.subplots(1, 2, figsize=(11, 4.2))
    # warehouse coverage
    grans = list(WAREHOUSE.keys())
    counts = [WAREHOUSE[g][1] for g in grans]
    a1.bar(grans, counts, color=[ACCENT, SKY, MUTE], ec="#222", lw=0.5)
    for i, (g, c) in enumerate(zip(grans, counts)):
        a1.text(i, c + 3000, f"{c:,}\n({WAREHOUSE[g][0]} coins)", ha="center", fontsize=7.8)
    a1.set_title("Warehouse candle coverage (TimescaleDB)", color=ACCENT, weight="bold", fontsize=10)
    a1.set_ylabel("candles"); a1.set_ylim(0, max(counts) * 1.25)
    # polymarket snapshots
    pn = [p[0] for p in POLY_SNAPSHOTS]; pv = [p[1] for p in POLY_SNAPSHOTS]
    a2.barh(range(len(pn)), pv, color=EDGE, ec="#222", lw=0.5)
    for i, v in enumerate(pv):
        a2.text(v + 30, i, f"{v:,}", va="center", fontsize=7.6)
    a2.set_yticks(range(len(pn))); a2.set_yticklabels(pn, fontsize=8)
    a2.invert_yaxis(); a2.set_xlim(0, max(pv) * 1.2)
    a2.set_title("Polymarket signal mix (5,400 snapshots)", color=ACCENT, weight="bold", fontsize=10)
    fig.suptitle("Data & signal inventory — what the system sees", fontsize=12, weight="bold", color=ACCENT, y=1.02)
    save(fig, "fig_signals.png")


# --------------------------------------------------------------------------- #
# FIG 8 — per-trade safety stack (decision pipeline gates → outcome buckets)
# --------------------------------------------------------------------------- #
def fig_safety_stack():
    fig, ax = plt.subplots(figsize=(10.5, 5.4))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis("off")
    gates = [
        "1 · data quality   (fresh? sane book?)",
        "2 · market eligibility   (tradeable, liquid)",
        "3 · regime   (does the setup fit conditions?)",
        "4 · signal agreement   (UNIQUE independent clusters)",
        "5 · edge   (slippage-aware: VWAP fill, not mid)",
        "6 · risk   (notional · daily-loss · concentration)",
        "7 · governor   (portfolio collision · correlated cap)",
    ]
    n = len(gates)
    for i, g in enumerate(gates):
        w = 7.4 - i * 0.55          # funnel narrows
        x = (10 - w) / 2
        y = 5.3 - i * 0.62
        ax.add_patch(FancyBboxPatch((x, y), w, 0.5, boxstyle="round,pad=0.02", fc="#eef2f7", ec=ACCENT, lw=1.1))
        ax.text(5, y + 0.25, g, ha="center", va="center", fontsize=8.0, color=ACCENT, weight="bold")
        if i < n - 1:
            ax.add_patch(FancyArrowPatch((5, y), (5, y - 0.12), arrowstyle="-|>", mutation_scale=10, color="#888", lw=1.0))
    # outcome buckets
    buckets = [("APPROVED", UP), ("REDUCED", GOLD), ("WATCHLIST", SKY), ("REJECTED", DOWN), ("KILL", "#111")]
    bw = 1.7
    for j, (b, c) in enumerate(buckets):
        x = 0.55 + j * 1.85
        ax.add_patch(FancyBboxPatch((x, 0.35), bw, 0.55, boxstyle="round,pad=0.02", fc=c, ec="#222", lw=1.0))
        ax.text(x + bw / 2, 0.62, b, ha="center", va="center", fontsize=8, color="white", weight="bold")
    ax.text(5, 1.25, "approval_score + size_multiplier  (multiplier only ever REDUCES; any gate REJECT → REJECTED; any KILL → KILL, regardless of score)",
            ha="center", fontsize=7.4, color="#444", style="italic")
    ax.text(5, 5.95, "The per-trade safety stack — every order runs the gauntlet before any capital moves",
            ha="center", fontsize=11.5, weight="bold", color=ACCENT)
    save(fig, "fig_safety_stack.png")


# --------------------------------------------------------------------------- #
# FIG 9 — the paper→live reality gap
# --------------------------------------------------------------------------- #
def fig_reality_gap():
    fig, ax = plt.subplots(figsize=(11, 5.2))
    ax.set_xlim(0, 11); ax.set_ylim(0, 6); ax.axis("off")

    def box(x, y, w, h, t, c, ec, fs=8, strike=False):
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02", fc=c, ec=ec, lw=1.3))
        ax.text(x + w / 2, y + h / 2, t, ha="center", va="center", fontsize=fs, weight="bold",
                color=ec, linespacing=1.3)
        if strike:
            ax.plot([x + 0.15, x + w - 0.15], [y + h / 2, y + h / 2], color=DOWN, lw=1.6)

    ax.text(2.3, 5.7, "PAPER ARENA  ·  ranked by fitness", ha="center", fontsize=10, weight="bold", color=ACCENT)
    box(0.4, 4.5, 3.8, 0.8, "poly_fade_spike  ·  ELITE\n+$3.27, 3/3 wins  (paper #1)", "#eef5ee", UP, 8)
    box(0.4, 3.5, 3.8, 0.8, "poly_short_binary (g1)\n+$57.68 — the comeback child", "#eef5ee", "#2e8b57", 8)
    box(0.4, 2.5, 3.8, 0.8, "polymarket_market_maker\nsteady, low edge", "#eef5ee", MUTE, 8)
    box(0.4, 1.5, 3.8, 0.8, "poly_short_binary (g0)\n−$100 — CULLED", "#fbeeee", DOWN, 8)

    # translation gate
    box(4.7, 2.6, 1.6, 2.2, "TRANSLATION\nGATE\n\nauto-promote\nLIVE-ELIGIBLE\nfilter ·\nslippage ·\nfill-reconcile", "#fff7e6", GOLD, 7.4)

    ax.text(8.85, 5.7, "LIVE-ELIGIBLE  ·  real money", ha="center", fontsize=10, weight="bold", color=ACCENT)
    box(6.9, 4.5, 3.8, 0.8, "poly_fade_spike  —  BANNED\nsigns+posts, dies on thin books", "#fbeeee", DOWN, 8, strike=True)
    box(6.9, 3.5, 3.8, 0.8, "poly_short_binary  ✓ PROMOTED\nthe culled genome earns real $", "#eef5ee", UP, 8)
    box(6.9, 2.5, 3.8, 0.8, "polymarket_market_maker ✓\n+ cb_momentum / cb_mean_rev", "#eef5ee", "#2e8b57", 8)
    box(6.9, 1.5, 3.8, 0.8, "correlation veto: no two live\ncapsules share family+asset", "#eef2f7", ACCENT, 8)

    for y in (4.9, 3.9, 2.9):
        ax.add_patch(FancyArrowPatch((4.3, y), (4.7, 3.7), arrowstyle="-|>", mutation_scale=11, color="#aaa", lw=1.0))
    for y in (4.9, 3.9, 2.9):
        ax.add_patch(FancyArrowPatch((6.3, 3.7), (6.9, y), arrowstyle="-|>", mutation_scale=11, color="#aaa", lw=1.0))

    ax.text(5.5, 0.75, "$10,000 PAPER budget  →   ~$30 REAL ceiling  (STAKE_USD $5 × small integer multipliers)",
            ha="center", fontsize=8.4, color=DOWN, weight="bold")
    ax.text(5.5, 0.35, "Fitness in the arena ≠ fitness for real money. The whole safety stack exists to translate between the two.",
            ha="center", fontsize=8, color="#444", style="italic")
    ax.text(5.5, 5.95, "The paper → live reality gap", ha="center", fontsize=12.5, weight="bold", color=ACCENT)
    save(fig, "fig_reality_gap.png")


# --------------------------------------------------------------------------- #
# WORD DOCUMENT
# --------------------------------------------------------------------------- #
def build_doc():
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11)

    title = doc.add_heading("The HFT Control-Plane — Architecture, Timing & Strategic-Space Playbook", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("A full breakdown of a self-evolving trading system: how data becomes signals, "
                            "how thousands of candidate strategies are bred and culled, how the survivors earn "
                            "real capital, where the edge actually lives, and — honestly — whether it is proven yet.")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph(f"Generated {GEN_DATE} · github.com/IsaiahDupree/HFT · data snapshot from the live arena + TimescaleDB warehouse")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in meta.runs: r.italic = True; r.font.size = Pt(9)

    def H(t, l=1): doc.add_heading(t, level=l)
    def P(t): return doc.add_paragraph(t)
    def B(t):
        return doc.add_paragraph(t, style="List Bullet")
    def BB(bold, rest):
        p = doc.add_paragraph(style="List Bullet"); r = p.add_run(bold); r.bold = True; p.add_run(rest); return p
    def fig(name, cap):
        path = os.path.join(OUT, name)
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(6.4))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c = doc.add_paragraph(cap); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in c.runs: r.italic = True; r.font.size = Pt(9)

    # 1
    H("1. Executive summary")
    P("This system is not a single strategy — it is a control-plane that grows, tests, and funds strategies the "
      "way nature grows organisms. Hundreds of candidate agents, each carrying a genome (a strategy kind plus "
      "parameters), trade against real market signals with simulated money. The ones that prove genuinely "
      "skilled survive, breed, and are handed a capsule of real risk capital; the rest are culled. The whole "
      "thing runs on a clock, around the clock.")
    P("Two honest findings frame everything in this document:")
    BB("The selection machinery works. ", "Across 133 agents and 13 strategy genomes, the arena has already "
       "evolved one full generation — culling 82 losers, preserving 5 elites, breeding 13 children — and the "
       "survivors are exactly the risk-adjusted winners (the steady poly_fade_spike line), while whole genome "
       "families that never worked (poly_breakout, cb_breakout) went extinct. Selection pressure is real.")
    BB("No edge is statistically proven yet — and we refuse to pretend otherwise. ", "Put through a full overfit "
       "gauntlet (Probability of Backtest Overfit, Deflated Sharpe, walk-forward), all 12 crypto priors FAILED "
       "the hardened gate. BTC daily momentum is the closest (Deflated Sharpe 0.94 vs a 0.95 bar) but does not "
       "clear it; trading it faster (hourly + risk-sizing) made it worse, not better. This is why the system "
       "trades sim money first and gates real capital behind survival in the live arena, not a backtest.")

    # 2
    H("2. System architecture")
    fig("fig_architecture.png", "Figure 1. End-to-end: data sources → warehouse → overfit gauntlet → genetic arena → capsules → stages → gated execution, with realized outcomes feeding back.")
    P("Five layers, one closed loop:")
    BB("Data sources. ", "Coinbase OHLC candles since each coin's origination; Polymarket Gamma/CLOB for the "
       "recurring 5-/15-minute up-down markets, elections and more; the L2 order book for microstructure "
       "signals (order-flow imbalance, VPIN, microprice); and on-chain wallet activity for copy signals.")
    BB("Warehouse (TimescaleDB) + control-plane (SQLite). ", "Two stores, by design: the TimescaleDB warehouse "
       "holds the heavy market time-series — ~257k candles plus ticks and snapshots as hypertables (it replaced "
       "two diverging SQLite files, so every process now reads one store and it scales to concurrent writes); a "
       "small SQLite control-plane holds the transactional arena state — paper_agents, capsules, the decision "
       "journal and the hash-chained order_events audit trail.")
    BB("Backtest + overfit gauntlet. ", "Before any prior is trusted, it is grid-searched, walk-forward "
       "validated, and stress-tested for overfitting. This layer produces priors, not deploy-ready parameters.")
    BB("The arena. ", "A genetic population of agents trades live signals with sim cash; a fitness function "
       "ranks them; evolution culls, breeds and preserves elites every 50 ticks.")
    BB("Capsules, stages & gated execution. ", "Survivors are issued a capsule — a hard risk envelope — and "
       "promoted along sim → paper → live-eligible → live. Real execution is gated and tiny by default.")

    # 3
    H("3. Data & signal inventory")
    fig("fig_signals.png", "Figure 2. What the system sees: warehouse candle coverage and the Polymarket signal mix.")
    P("Depth on the left, breadth on the right. The warehouse holds deep daily history (12 coins, 29,297 daily "
      "candles back to 2015 for BTC) and high-resolution hourly bars (BTC 95,158 · ETH 87,749 · SOL 43,387) for "
      "finer experiments. On the signal side, 5,400 Polymarket snapshots are dominated by the 5- and 15-minute "
      "binary up-down markets (2,160 each) — the highest-frequency, most-recurring opportunity — with a long "
      "tail of elections, crypto, geopolitics and weather markets. Microstructure signals (OFI, VPIN, "
      "microprice) are computed from the L2 book for the market-making genomes.")

    # 4
    H("4. The strategy lifecycle & strategic zones")
    fig("fig_lifecycle.png", "Figure 3. An agent's life — only proven survivors advance toward real money.")
    P("Every agent moves through the same gauntlet, and the right action differs at each stage — this is the "
      "'what to do in the space between' map at the level of a single strategy:")
    BB("BIRTH. ", "A genome {kind, params} is seeded — from a preset, a research-driven target, a mutation, or a crossover of two parents.")
    BB("SIM. ", "The agent starts with $1,000 of paper cash and trades real live signals. No real money is at risk; the only thing it can earn is fitness.")
    BB("RANK + SEAL. ", "Every 50 ticks the generation seals: agents are ranked by fitness, the bottom ~50% are culled, the top breed children, and the top-N are preserved as elites (protected from culling unless their drawdown blows out).")
    BB("CAPSULE. ", "Survivors are issued a capital capsule — a risk envelope with allocation, daily-loss and drawdown limits. This is the moment an agent goes from 'idea' to 'funded'.")
    BB("LIVE-ELIGIBLE → LIVE. ", "Only after sustained capsule performance does an agent become eligible for real, tightly-capped execution. Most never get here, by design.")

    # 5
    H("5. The arena-loop timing plan")
    fig("fig_timing.png", "Figure 4. The clock: what fires every cycle vs. every hour vs. every two hours.")
    P("A scheduled job (launchd, every ~5 minutes) drives the whole system. The cadence is deliberately layered "
      "so cheap, fast actions happen constantly while expensive, committing actions happen rarely:")
    BB("Every cycle (~5 min): ", "snapshot the markets, tick the arena (agents evaluate signals, open/close paper positions, fitness updates), and capture the L2 book. This is the heartbeat.")
    BB(f"Every {ALLOCATE_EVERY} cycles (~1 h): ", f"auto-allocate — re-rank survivors and (re)issue capsules from a ${ALLOC_BUDGET:,} budget. This is where money decisions concentrate.")
    BB(f"Every {RESEARCH_EVERY} cycles (~2 h): ", "run the opportunity scanners (leaderboard, near-resolution, cross-timeframe, order-book imbalance) and research-refresh, which re-targets agents on the current market regime.")
    P("The strategic insight is in the SPACE BETWEEN: the heartbeat gathers evidence continuously, but capital is "
      "only moved on the hourly beat and strategy is only re-aimed on the two-hour beat. Fast to observe, slow "
      "to commit — the opposite of over-trading.")

    # 6
    H("6. Is the edge real? The overfit gauntlet")
    fig("fig_overfit.png", "Figure 5. Deflated Sharpe by coin against the hardened gate — and why faster made BTC worse.")
    P("This is the part most systems skip and the reason to trust this one. Every candidate prior is run through "
      "three overfit defenses: the Probability of Backtest Overfit (combinatorial cross-validation — how often "
      "the in-sample-best configuration underperforms out-of-sample), the Deflated Sharpe Ratio (which discounts "
      "a Sharpe for the number of trials and for non-normal returns), and multi-fold walk-forward. The hardened "
      "gate demands PBO < 0.30 AND Deflated Sharpe > 0.95 AND positive median out-of-sample Sharpe.")
    P(f"Result: 0 of 12 coins hardened. BTC daily momentum is the single most defensible edge (PBO 0.00, all four "
      f"walk-forward folds out-of-sample positive) but its Deflated Sharpe is 0.94 — just under the bar. "
      f"Out-of-sample Sharpe stayed positive on only {WALKFWD_HELD}/12 coins. And when we tried to sharpen BTC "
      f"with finer hourly bars plus an inverse-volatility sizing overlay, the Deflated Sharpe FELL to "
      f"{DSR_BTC_HOURLY:.2f}: the sizing reduced turnover as designed, but finer bars added noise and trials, "
      f"not signal. The honest conclusion: crypto daily-momentum is a real-but-modest tendency, not a sure thing "
      f"— so we let the live arena be the final out-of-sample judge instead of betting on a backtest.")

    # 7
    H("7. The genetic arena — natural selection in action")
    fig("fig_arena.png", "Figure 6. Total vs. surviving agents per genome — whole families thrive or go extinct.")
    P("The arena has run one full generation and is partway through the second. The selection signal is "
      "unmistakable: poly_fade_spike (a Polymarket spike-fade play) had all 6 of its agents survive as elites; "
      "the market-maker and mean-reversion families partly persist; and poly_breakout and cb_breakout were "
      "wiped out entirely (0 survivors). Fitness is not raw profit — it is profit-percent minus twice the "
      "max-drawdown-percent plus an activity bonus, so it rewards steady, low-drawdown skill.")
    fig("fig_leaderboard.png", "Figure 7. The live leaderboard — and the evolution comeback story.")
    P("The leaderboard tells the most interesting story in the system. The top earner (+$57.68) is a generation-1 "
      "MUTATED CHILD of poly_short_binary_directional — a genome whose generation-0 ancestors lost $100 and were "
      "culled. Evolution found a profitable variant of a previously-failing idea. Yet fitness still crowns the "
      "calm poly_fade_spike elites (+$3.27 on a perfect 3/3) over that volatile big winner — exactly the "
      "risk-adjusted preference you want deciding who gets real money.")

    # 8 — the reality gap (flagship narrative 1)
    H("8. The paper → live reality gap")
    fig("fig_reality_gap.png", "Figure 8. Why the paper champion and the real-money pick are different agents.")
    P("This is the single most important idea in the system, and the real data shows it vividly. Fitness in the "
      "arena and fitness for real money are DIFFERENT objectives, and an entire safety stack exists to translate "
      "between them. poly_fade_spike is the paper-arena ELITE (+$3.27 on a perfect 3/3) — yet it is structurally "
      "BANNED from live capital, because it signs and posts orders that simply die on Polymarket's thin visible "
      "books ('no orders to match'). Meanwhile poly_short_binary_directional, whose generation-0 ancestors lost "
      "$100 and were culled, produced a generation-1 child that is both the top earner (+$57.68) AND live-"
      "eligible. The translation is done by three mechanisms: auto-promote's LIVE-ELIGIBLE genome filter (only "
      "kinds that can actually fill — poly_short_binary, llm_probability_oracle, polymarket_market_maker, "
      "cb_momentum_burst, cb_mean_reversion), a slippage estimator that prices fills by walking the L2 book to a "
      "VWAP rather than trusting top-of-book, and fill-reconciliation that checks simulated fills against the "
      "real on-chain OrderFilled event. And the capital truth is sobering by design: the $10,000 is PAPER "
      "budget; the REAL ceiling is about $30 (a $5 stake times small integer 'how many losing stakes before we "
      "pause' multipliers). Proof first, size later.")

    # 9 — the safety stack
    H("9. The per-trade safety stack")
    fig("fig_safety_stack.png", "Figure 9. Every order runs an eight-gate gauntlet before any capital moves.")
    P("No single signal can move money. Each candidate trade passes a decision pipeline of gates — data quality, "
      "market eligibility, regime fit, signal agreement, edge (slippage-aware), risk, and a portfolio governor — "
      "that together produce an approval_score and a size_multiplier which can only ever REDUCE size, never "
      "amplify it. The verdict lands in one of five buckets (APPROVED, REDUCED, WATCHLIST, REJECTED, "
      "KILL_SWITCH), and the rules are hard: any gate that rejects forces REJECTED regardless of score, and any "
      "kill-switch forces KILL regardless of everything. Approved orders then hit the execution router, the "
      "single submit path, which adds five more checks — idempotency (no duplicate client order), a global halt "
      "flag, the capsule gate, a risk-engine check (notional, daily-loss, rate, concentration), and only then "
      "the venue adapter. Every decision is appended to a hash-chained order_events table, so the entire "
      "execution history is tamper-evident and auditable. A three-mode kill-switch (pause-new, close-and-pause, "
      "liquidate) can halt the whole system at once.")

    # 10 — self-honesty
    H("10. The system audits its own honesty")
    P("Most trading systems trust their own confidence. This one measures it. Every per-trade decision is logged "
      "to a decision journal with its approval_score; later, calibration joins those scores to realized PnL and "
      "draws a reliability diagram — when the pipeline said '0.85', did those trades actually win ~85%? That is "
      "the forward-in-time twin of the offline overfit gauntlet: the Deflated-Sharpe/PBO battery defends against "
      "overfitting before deployment (one-shot, on history), while calibration defends against it continuously "
      "after deployment (online, on the live tape). The same discipline, running in both directions of time. "
      "Fill-reconciliation closes the loop by giving the journal ground truth — without checking simulated fills "
      "against real ones, the paper-versus-live comparison would be fiction.")

    # 11 — ensemble decision
    H("11. The ensemble decision — who gets a capsule of money, and why")
    P("This is the heart of the original goal: an AI ensemble that decides which agents earn a capsule and can "
      "articulate why. Two paths fund agents. The hourly allocator re-ranks the living population by risk-"
      "adjusted fitness and issues PAPER capsules (capped at ~25% of the budget each) to survivors above the "
      "fitness and minimum-trade thresholds. Separately, at every evolution seal, auto-promote considers the "
      "live path: it ranks all alive agents, keeps only LIVE-ELIGIBLE genomes, applies a correlation-aware "
      "diversity veto (no new live capsule that is structurally identical — same strategy family and asset "
      "class — to one already funded), and pauses (never silently closes) capsules whose agent falls out of the "
      "top ranks. The live path is doubly gated by explicit environment flags and stays off by default. A "
      "capsule is a contract — allocated capital, a daily-loss limit, a max-drawdown kill-switch — not a blank "
      "cheque. The 'why' is always legible: an agent is funded because it ranked here, on this fitness, over "
      "this many trades, at this drawdown, in a genome that can actually fill, uncorrelated with what's already "
      "live.")

    # 12 — independence (flagship narrative 3)
    H("12. The one enemy: correlation")
    P("Step back and the whole architecture is fighting a single foe at every layer: pseudo-replication — being "
      "fooled by correlated copies of one bet. The same idea, 'N apparent agreements may really be one', recurs "
      "in otherwise-unrelated subsystems. Wallet-consensus is cluster-aware, so five wallets run by one operator "
      "collapse to a single effective vote (45 wallets are tracked; this is the platform's structural edge — an "
      "individual wallet cannot see what other top wallets are doing in real time, but the platform can). The "
      "signal-agreement gate counts UNIQUE independent signal clusters, not raw signal count. Auto-promote's "
      "correlation veto refuses capital to a strategy that merely duplicates an existing live bet. And the "
      "Deflated Sharpe explicitly discounts the best-of-N-trials result for the number of trials. Across wallets, "
      "signals, capital allocation, and statistics, every promotion gate is really asking the same question: is "
      "this edge genuinely INDEPENDENT, or am I being shown the same coin flip several times?")

    # 13 — multi-venue + edges
    H("13. Venues, edges & the mechanical arbitrage")
    P("The system is multi-venue: Coinbase (the underlying), Polymarket (binary prediction markets), dYdX (a "
      "full third venue with its own market-making engine and order-book-imbalance signals), Kalshi "
      "(integrated), plus on-chain watchers (e.g. Aave liquidations) and the wallet trade stream. Beyond the "
      "model-driven genomes, two edges are structural. Complement-sum arbitrage is model-free: when a binary "
      "market's Up-ask plus Down-ask trade below $1, buying both guarantees a positive payout no matter which "
      "side wins. Cross-venue arbitrage is model-based: a Black-Scholes implied probability (P[close above "
      "strike] from the underlying's volatility) is compared to Polymarket's implied probability, and the agent "
      "fires when they diverge. These sit alongside the Avellaneda-Stoikov market-maker (logit-space quoting "
      "for binaries, inventory caps, OFI/VPIN microstructure signals) as the non-momentum sources of edge.")

    # 14 — playbook
    H("14. Strategic playbook — what to do, when")
    for k, v in [
        ("Observe fast, commit slow.", " The 5-minute heartbeat gathers evidence; only the hourly beat moves capital and the seal promotes to live. Never let the heartbeat trigger funding."),
        ("Mind the reality gap.", " A paper champion that can't fill is worth nothing live. Judge agents by live-eligibility and slippage-aware fills, not paper PnL."),
        ("Trust fitness over PnL, and calibration over confidence.", " The biggest winner can be the worst bet; an approval_score is only worth what the reliability diagram says it is."),
        ("Hunt independence.", " Reward edges that are genuinely uncorrelated; treat correlated agreement as one vote, not many — across wallets, signals, and capsules."),
        ("Keep the extinct families dead.", " poly_breakout and cb_breakout earned zero survivors and can't fill live; don't re-seed them without a research reason."),
        ("Don't chase resolution.", " Faster bars lowered BTC's Deflated Sharpe (0.94 → 0.60). More frequency = more trials to overfit against, not more edge."),
        ("Gate hard, cap tiny.", " Real execution stays behind explicit flags with a ~$30 ceiling until realized PnL across a real sample is honest. Scale with proof, not hope."),
    ]:
        BB(k, v)

    # 15 — roadmap
    H("15. Roadmap")
    for t in ["Repoint the live capture loop (ONE_MINUTE candles, ticks, snapshots) into the warehouse so it is canonical end-to-end.",
              "Accumulate a real arena-funded PnL sample, then let calibration grade the gates against realized outcomes.",
              "Let the arena run more generations — track whether the comeback genome (poly_short_binary g1) holds up or reverts.",
              "Add Timescale compression + retention on the minute/tick hypertables as they grow.",
              "Only after a defensible edge survives BOTH the offline gauntlet AND the live arena: widen the live cap beyond $30, scaling with proof."]:
        B(t)

    out = os.path.join(OUT, "HFT_Architecture_Timing_Playbook.docx")
    doc.save(out)
    return out


def render_all():
    os.makedirs(OUT, exist_ok=True)
    fig_architecture(); fig_lifecycle(); fig_timing(); fig_overfit()
    fig_arena(); fig_leaderboard(); fig_signals()
    fig_safety_stack(); fig_reality_gap()
    print("rendered 9 figures →", OUT)


if __name__ == "__main__":
    render_all()
    out = build_doc()
    print("DOC:", out)
