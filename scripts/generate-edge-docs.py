#!/usr/bin/env python3
"""
generate-edge-docs — render every edge in the catalog (docs/EDGES.md is the
source of truth) as a Word document under docs/edges/. One .docx per live
edge + 00-INDEX + the graveyard of rejected edges. Idempotent: rerun after
any verdict changes. Requires python-docx.

    python3 scripts/generate-edge-docs.py
"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent.parent / "docs" / "edges"
TODAY = "2026-06-11"

STATUS_COLORS = {
    "GO": RGBColor(0x1B, 0x7A, 0x2B),
    "PAPER": RGBColor(0xB0, 0x7A, 0x00),
    "NEGATIVE": RGBColor(0xB0, 0x30, 0x20),
    "REJECTED": RGBColor(0x80, 0x80, 0x80),
}


def add_doc(filename, title, status, status_kind, sections):
    """sections: list of (heading, body) where body is str | list[str] (bullets)
    | ('table', [[...], ...]) with row 0 as header."""
    d = Document()
    style = d.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    h = d.add_heading(title, level=0)
    for run in h.runs:
        run.font.size = Pt(20)

    p = d.add_paragraph()
    r = p.add_run(f"STATUS: {status}")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = STATUS_COLORS[status_kind]
    meta = d.add_paragraph()
    mr = meta.add_run(
        f"As of {TODAY} · catalog of record: HFT-work/docs/EDGES.md · "
        "stance: TRADING_POLICY.md (pro-trading, anti-delusion — the only veto is the data)"
    )
    mr.italic = True
    mr.font.size = Pt(8.5)

    for heading, body in sections:
        d.add_heading(heading, level=1)
        if isinstance(body, str):
            d.add_paragraph(body)
        elif isinstance(body, list):
            for item in body:
                d.add_paragraph(item, style="List Bullet")
        elif isinstance(body, tuple) and body[0] == "table":
            rows = body[1]
            t = d.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Light Grid Accent 1"
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    c = t.cell(i, j)
                    c.text = str(cell)
                    for para in c.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(9.5)
                            if i == 0:
                                run.bold = True

    foot = d.add_paragraph()
    fr = foot.add_run(
        "Never claim guaranteed profit · never martingale · forward paper-track with "
        "independent resolution beats any in-sample number."
    )
    fr.italic = True
    fr.font.size = Pt(8)
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.mkdir(parents=True, exist_ok=True)
    d.save(OUT / filename)
    print(f"  wrote {filename}")


EDGES = []  # (filename, title, status, kind, sections) — appended below

EDGES.append((
    "01-funding-carry-persistent-alts.docx",
    "Edge #1 — Funding Carry on Persistent-Funding Alts",
    "LEANS GO (paper → small live)", "GO",
    [
        ("The trade",
         "Delta-neutral carry: short the perp + long spot on alts whose funding is persistently "
         "positive, harvesting funding 3×/day while the price legs cancel. The edge is the funding "
         "income; the risks are the spot-perp basis and execution cost. Universe selection is the "
         "whole game: persistence-ranked names (sign-stability × magnitude ÷ flips) have low "
         "turnover, so entry fees are paid rarely while funding is collected many times."),
        ("Evidence (full gauntlet)", ("table", [
            ["Test", "Result"],
            ["Lake backtest, 28 high-funding alts, 500d", "every variant positive; best +22% APR, Sharpe 8.07"],
            ["Walk-forward OOS", "held (OOS Sharpe 7.51)"],
            ["PBO / DSR", "0.00 / 1.00 — not overfit"],
            ["Basis-aware re-test (real spot+perp)", "Sharpe 10.8 → 6.7 (~38% haircut) — real but survives"],
            ["Data-measured fee (20% L2 fill → 4.22bp/leg)", "+37% APR, Sharpe 7.7 — survives 4× realistic fees"],
            ["Advisor", "BUY (92/100)"],
        ])),
        ("Caveats & falsifiers", [
            "Gate on DURABLE (median) funding, never the mean — spike-inflated means are uncollectable (AZTEC mean +26% vs median +11%).",
            "Fill-rate calibrated on liquid dYdX majors → optimistic for the alts actually traded.",
            "Basis blowout in a squeeze is the tail risk; borrow availability is why funding is high.",
            "Execution wall (2026-06-05): 0 carries executable on Hyperliquid (no spot borrow; +11%/hr funding floor → +8% net, under the bar). Needs Binance-global proxy or dYdX legs.",
            "Stays paper until a name clears durable funding > ~+16% APR with a deep same-asset hedge.",
        ]),
        ("Tooling", "discover:high-funding → fetch:funding:binance → discover:funding-persistence → "
                    "backtest:carry-neutral → backtest:basis → carry:maker-fill · forward: carry:monitor (hourly cron)."),
    ],
))

EDGES.append((
    "02-calendar-basis-carry.docx",
    "Edge #2 — Calendar (Dated-Futures) Basis Carry",
    "PAPER → DEPLOYABLE (the cleanest edge in the book)", "GO",
    [
        ("The trade",
         "Cash-and-carry: long spot + short the front-quarter dated future in contango. At delivery "
         "the future MUST converge to spot, so the entry basis is harvested with the convergence "
         "locked — no funding cash-flows, no open-ended basis risk if held to expiry."),
        ("Evidence", ("table", [
            ["Test", "Result"],
            ["Annualized basis", "BTC +8.07% (contango 100% of days), ETH +7.66% (98%)"],
            ["Best variant", "Sharpe 3.06, +8.5%/yr, +24.7% over 999d"],
            ["PBO / DSR", "0.00 / 1.00"],
            ["Carry signature", "realized return == observed basis (true carry, not artifact)"],
            ["No-lookahead", "verified; contract-roll seams skipped"],
            ["Advisor", "PAPER (72) — clean metrics, modest absolute return"],
        ])),
        ("Caveats & falsifiers", [
            "~8%/yr is modest; daily MTM is volatile even though the terminal payoff is locked.",
            "Roll/execution costs + margin on the short-future leg eat into the headline.",
            "Needs a dated-futures venue (Binance/Deribit) — Hyperliquid lacks one.",
            "Term-structure timing variants are falsified — full-life hold is best; don't get clever.",
        ]),
        ("Tooling", "npm run backtest:calendar-basis · core math: calendarBasisReturns (pure, tested)."),
    ],
))

EDGES.append((
    "03-vol-risk-premium.docx",
    "Edge #3 — Volatility Risk Premium (Sell Vol)",
    "PAPER (real but tail-risky — size for the tail, not the Sharpe)", "PAPER",
    [
        ("The trade",
         "Implied vol systematically exceeds subsequently-realized vol; option sellers are paid for "
         "bearing variance risk. Sell straddles/variance on BTC and harvest the gap."),
        ("Evidence (Deribit DVOL vs Binance realized, 2021–2026)", ("table", [
            ["Test", "Result"],
            ["VRP", "+8.87 vol points; positive 73% of 1,867 days"],
            ["Honest non-overlapping Sharpe", "≈1.23 (36 independent 30d blocks, won 24/36)"],
            ["PBO / DSR", "0.17 / 1.00"],
            ["Overlapping-window Sharpe", "9.26 — inflated, correctly DISCARDED"],
        ])),
        ("Caveats & falsifiers", [
            "Fat left tail / negative skew: small premiums, occasional large losses in vol spikes — 'pennies in front of a steamroller'.",
            "Sharpe does not capture the skew; the honest 1.23 still overstates risk-adjusted appeal.",
            "Position sizing must assume the tail event happens while positioned.",
        ]),
        ("Tooling", "scripts/_carry-deribit-vol-risk-premium.ts · vol-surface forward track live (skew + term + VRP)."),
    ],
))

EDGES.append((
    "04-staking-hedged-yield.docx",
    "Adjacent — Staking-Hedged Yield Carry",
    "PAPER (real structural carry, inflated headline)", "PAPER",
    [
        ("The trade",
         "Stake ETH/SOL (~3.2%/7% APY), short the perp to hedge price → delta-neutral staking yield; "
         "the short also collects positive funding when present. ~5–6% net APR, fee-robust "
         "(survives 150 bps/yr drag)."),
        ("Caveats & falsifiers", [
            "Headline Sharpe (~14) is inflated: models staking as riskless.",
            "ETH unbond-queue illiquidity: can't unwind the hedge against a leg you can't exit for days/weeks.",
            "Slashing, LST depeg, tracking error are real and unmodeled.",
        ]),
        ("Tooling", "scripts/_carry-staking-hedged-yield.ts"),
    ],
))

EDGES.append((
    "05-hl-delta-funding-hold.docx",
    "HL-Delta — Same-Venue Funding HOLD Sleeve",
    "PAPER SLEEVE (the HOLD is real; the rotation is a cost-line loss)", "PAPER",
    [
        ("The trade",
         "Hold the delta-neutral funding-receiving position on a single persistent name on "
         "Hyperliquid (HYPE), collecting same-venue funding without cross-venue legs. Mined from "
         "the strategy fan-out and pushed through the full gauntlet."),
        ("Evidence", ("table", [
            ["Test", "Result"],
            ["OOS (HYPE hold)", "+7.4% APR out-of-sample"],
            ["Rotation variant", "FAILS — rotating names loses the edge to costs/timing"],
            ["Front gate", "passed the ported Freqtrade lookahead + recursive bias detectors"],
        ])),
        ("Caveats & falsifiers", [
            "Single-name concentration: the sleeve IS HYPE; if its funding regime ends, the sleeve ends.",
            "Same Hyperliquid execution constraints as Edge #1 (funding floor, no spot borrow for negatives).",
        ]),
        ("Tooling", "see commits 13f7636 / b15498b (hl-delta gauntlet + verdict)."),
    ],
))

EDGES.append((
    "06-polymarket-merge-maker.docx",
    "Polymarket — Merge-Maker + Rebates (Lane A)",
    "PAPER, FORWARD-ACCRUING — the strongest prediction-market lane", "PAPER",
    [
        ("The trade",
         "On short-duration crypto binaries (5-min/15-min/hourly Up/Down): post maker BUY bids on "
         "BOTH the Yes and No tokens such that the paired cost is below $1.00, merge complete sets "
         "back to $1 USDC (riskless margin at merge), collect maker rebates. Inventory risk nets "
         "out at the pair level — the structural answer to the adverse selection that kills naked "
         "quoting (see doc 07)."),
        ("Existence proofs (on-chain, verified 2026-06-10/11)", ("table", [
            ["Wallet", "Evidence"],
            ["coinman2 (0x55be…dca3)", "+$1,090,424 / 558d, ~1.4% of $79.1M volume; MAKER_REBATE + MERGE fingerprint; survived −$349k max DD"],
            ["0x6db568e6…", "+$1.58M, 7.2% margin, 54d — merge-maker hybrid, live today"],
            ["Bonereaper", "+$999k / 77d, 0.69% margin on volume, max DD only −$8.7k — the benchmark"],
        ])),
        ("Calibration fingerprint (what the profitable cohort actually does)", [
            "$1–14 median fills (the −$1.29M dead maker ran 4× bigger clips).",
            "Quote the FULL 0.02→0.99 lifecycle — only 5–11% of fills near the mid.",
            "~3s requote cadence (dead maker: 13s).",
            "Net 0.1–0.7% of volume — profit comes from volume × tiny margin, not from being right.",
        ]),
        ("Our forward paper (TradingBot2 Lane A)", [
            "merge_maker.py daemon on 5m/15m updowns, 120s cycles; cross-cycle pair state persisted.",
            "First 4 finalized windows: 73.7% pair completion; income decomposition MAKER-DRIVEN (+$24.10 maker vs +$3.50 residual). Tiny n — let it accrue.",
            "Independent prior: evan-kolberg's forward-validated passive_pair_accumulation (≈ the same strategy).",
        ]),
        ("Falsifiers / gates", [
            "Pair-completion rate collapsing → unpaired inventory is naked directional risk.",
            "G3 backtest on PMXT free L2 archive (queue-position fills) before any capital.",
            "Bankroll honesty: coinman2 needed six figures + $349k DD tolerance for $2k/day; small bankroll ⇒ proportionally small absolute PnL.",
        ]),
    ],
))

EDGES.append((
    "07-polymarket-binary-fv-maker.docx",
    "Polymarket — Binary Fair-Value Naked Maker (G2 Testbed)",
    "PAPER — CURRENTLY NEGATIVE; do not promote (testbed for the merge lane)", "NEGATIVE",
    [
        ("The trade (as built)",
         "Two-sided quotes on a single binary token around an independent CEX-feed fair value "
         "(digital-option pricer + A-S inventory skew), trade-driven paper fills, 24/7 daemon "
         "rolling market-to-market across the 5-min/15-min/hourly Up/Down series."),
        ("Honest forward readout (first ~18h of clean, correct-strike data)", ("table", [
            ["Metric", "Value"],
            ["Sessions / fills", "183 / 1,650"],
            ["Paper PnL (optimistic front-of-queue fills)", "−$2,066"],
            ["Rebates", "+$118"],
            ["Model A/B (13,566 ticks)", "baseline Brier 0.1068 BEATS enhanced 0.1105 — momentum+Hurst not earning its keep"],
            ["MARKET MID benchmark (the make-or-break)", "mid Brier 0.1063 BEATS both models — the fair value has NO edge over the book; the maker was supplying edge, not capturing it"],
            ["Diagnosis", "naked inventory rides directional moves into resolution; adverse selection ≫ spread+rebates; and the fair the quotes anchor on is no better than the mid"],
        ])),
        ("What it taught (why it stays running)", [
            "FORMAL DEMOTION (2026-06-11): naked quoting off this fair value is dead — promotion requires a fair value that beats the mid's Brier, a bar now wired into maker:paper:report.",
            "Every profitable wallet examined pairs/merges; none holds naked single-token inventory — the lane's capital path converges to the merge-maker (doc 06).",
            "The wrong-strike bug (5-min series priced off the 15-min candle) was caught BY the divergence self-check — the honesty machinery works.",
        ]),
        ("Infrastructure (reusable regardless of verdict)", [
            "Live A/B logging (p_fair_base/p_fair_enh/hurst per tick), strike-sanity abort, duration-scaled estimator windows, question forensics, knife-edge-excluded Brier scoring.",
        ]),
    ],
))

EDGES.append((
    "08-polymarket-copy-0x418d51e1.docx",
    "Polymarket — Copy Lane: 0x418d51e1 (sole survivor of 148)",
    "PAPER — copyable at ≤300s; 2-week forward shadow gates any sizing", "PAPER",
    [
        ("The trade",
         "Mirror wallet 0x418d51e1's NBA/MLB moneyline value entries (~1 trade/day) with up to "
         "5 minutes of delay. Lag-aware backtest (1-min CLOB history, +1¢ spread crossing, no "
         "lookahead): copy +13.5% ROI at BOTH 60s and 300s vs the leader's +14.8% on the same "
         "103 bets — retains 91% of edge because the alpha is SIDE SELECTION, not entry timing."),
        ("Why it survives when bigger names died", ("table", [
            ["Wallet", "Copy verdict"],
            ["0x418d51e1 (n=103)", "COPYABLE ≤300s: +13.5% vs +14.8%; slippage 0.5¢, drift negative (value entries)"],
            ["alwaysfade (n=19 ≥$1k clips)", "conditional: +10.1% @60s, delay-flat — n too thin, shadow only"],
            ["ethanaz (n=944)", "DEMOTED: copy −2.6% @60s vs leader +5.3% — post-fill drift IS his alpha; copier pays it"],
        ])),
        ("Skeptic / advocate checks", [
            "Shuffle (random timing, same sides): +14.8% → alpha is side selection (precisely what makes it copyable).",
            "Beta baseline (buy the favorite at the same times): −0.6% — not favorite-beta.",
            "Open flags: possible insider information channel; n=103; 5s/30s delays unmeasured (tape cap).",
        ]),
        ("Gate before any capital", [
            "Extend observe:wallet to record detection latency, live ask/spread at +60s/+300s, and leader clip size.",
            "2-week forward shadow with independent resolution; verdict settles on shadow data only.",
        ]),
    ],
))

EDGES.append((
    "09-polymarket-lp-rewards.docx",
    "Polymarket — Liquidity-Rewards Lane",
    "PAPER (real income, heavy-tail variance and concentration)", "PAPER",
    [
        ("The trade",
         "Rest quotes inside reward bands on Polymarket markets to farm the liquidity-rewards "
         "program; income = rewards + (incidental) inventory PnL from fills."),
        ("Honest ledger (TradingBot2, Jun 5–10, after the inventory-accounting fix)", ("table", [
            ["Metric", "Value"],
            ["Rewards", "+$630"],
            ["Inventory PnL (path-cross measured, capital-capped)", "+$119"],
            ["Total ≈", "+$749 ≈ $134/day on a $200 paper stake"],
            ["Concentration", "170/191 pools zero-fill — income comes from a few pools"],
            ["Variance", "inventory term ±$400 heavy-tail"],
        ])),
        ("Falsifiers / notes", [
            "The first ledger wrote inventory_pnl = 0.0 blind — the +$134/day only became credible after three accounting fixes; keep auditing the ledger, not the headline.",
            "Reward-program parameters can change at Polymarket's discretion — platform risk.",
            "Judge after weeks of accrual; size for the inventory tail.",
        ]),
    ],
))

GRAVEYARD_ROWS = [
    ["Edge", "Verdict", "The killing evidence"],
    ["Polymarket taker latency-snipe (5-min)", "LANE EMPTY", "0.5s median lag across 39 cron measurements; 0 stale events; 0 profitable takers among 3,189 wallets / 33k fills"],
    ["Naive copy-trading", "FORWARD-FALSIFIED ×2", "21-daemon forward walk + lag-aware backtests; ethanaz copy −2.6% @60s vs leader +5.3%"],
    ["Near-cert 95–99¢ grinding", "DEBUNKED as structural", "long-span grinders net LOSE (richerZ −$94k/514d); wins are selection-driven"],
    ["Stablecoin peg mean-reversion (continuous)", "REJECTED", "calm-regime clearly negative; it's a rare-EVENT trade, not a continuous edge"],
    ["Liquidation-cascade reversal strategy", "SHELVED", "failed its 5-year backtest honestly"],
    ["3-down reversion", "FORWARD-FALSIFIED", "TradingBot2 forward walk"],
    ["HL funding rotation (vs hold)", "FAILS", "rotation loses the HYPE hold's +7.4% OOS to costs/timing"],
    ["Momentum/RS rotation", "HOLD_BETA", "underperforms buy-and-hold; shuffle p=1.000"],
    ["Pairs / stat-arb (daily)", "HOLD_BETA", "BH +16,068% vs pairs −58%"],
    ["Minute pairs stat-arb", "REAL GROSS, DEAD NET", "+4bp/trade gross Sharpe ~4.5 < 24bp two-leg cost"],
    ["Cross-sectional funding carry", "REJECTED", "funding-only +8.9 Sharpe → price-aware −65%; loose basket = directional bet in disguise"],
    ["Intraday families (cascade/ignition/vol-spike/x-exch basis)", "ALL REJECTED", "real 2–10bp gross inefficiencies < 12–24bp retail costs"],
    ["Funding time-of-day / OI-squeeze / breadth timing / vol-reversion", "REJECTED", "controls: identical slots / shuffle p=0.92 / net-negative OOS / vol persists"],
    ["Negative-funding carry via DeFi borrow", "STRUCTURALLY CLOSED", "borrowable ∩ fat-negative-funding = ∅ on Aave V3"],
]

EDGES.append((
    "10-edge-graveyard.docx",
    "The Edge Graveyard — Tested and Rejected",
    "REJECTED / FALSIFIED — kept on record so they are not re-derived", "REJECTED",
    [
        ("Why this document exists",
         "Every entry below consumed real research effort and DIED at a control: a shuffle test, a "
         "beta baseline, honest cost accounting, a forward paper-track, or wallet forensics. The "
         "graveyard is the proof the gauntlet works — and the guard against re-deriving a dead idea "
         "the next time it looks shiny. Failing to find edge ≠ edge doesn't exist; but a control "
         "that fired IS a verdict."),
        ("The graveyard", ("table", GRAVEYARD_ROWS)),
        ("The two meta-lessons", [
            "Carry is real only when the hedge is TIGHT (own-spot, locked convergence); loose-basket carry is a directional bet in disguise.",
            "Intraday gross inefficiencies are real but smaller than retail round-trip costs — the harvesters are the fee-advantaged (maker rebates, colocation). Become the maker or stand aside.",
        ]),
    ],
))


def build_index():
    rows = [["Doc", "Edge", "Status"]]
    for fn, title, status, kind, _ in EDGES:
        rows.append([fn.split("-")[0], title, status])
    add_doc(
        "00-INDEX.docx",
        "Edge Catalog — Index",
        f"{len(EDGES)} documents · source of truth: docs/EDGES.md", "PAPER",
        [
            ("How to read this catalog",
             "One document per edge. Statuses: GO/leans-GO (cleared the gauntlet, deployment-path "
             "defined) · PAPER (real signal, forward track or gate pending) · NEGATIVE (honest "
             "losing readout, kept as testbed) · REJECTED (a control fired). The gauntlet every "
             "edge faces: no-lookahead lake backtest → walk-forward OOS → PBO + Deflated Sharpe → "
             "shuffle/permutation control → one-voice advisor → execution realism (data-measured "
             "fees, basis). A forward paper-track with independent resolution outranks every "
             "in-sample number."),
            ("Contents", ("table", rows)),
        ],
    )


if __name__ == "__main__":
    print(f"generating {len(EDGES)} edge docs + index → {OUT}")
    for args in EDGES:
        add_doc(*args)
    build_index()
    print("done.")
